import time
import pyautogui
import docx as aw
from docx.oxml.ns import qn
import os
from aims import aims

# Define Aims
aimList = aims
# Define Doc Details
heading = 'Practical - 3'
aim = 'To implement Caesar cipher encryption-decryption.'
headingSize = 22
aimSize = 14

# Create a new document
doc = aw.Document()

# Set Document Margins to normal
sections = doc.sections
for section in sections:
    section.top_margin = aw.shared.Inches(1)
    section.bottom_margin = aw.shared.Inches(1)
    section.left_margin = aw.shared.Inches(1)
    section.right_margin = aw.shared.Inches(1)

def makedoc(heading,aim,headingSize,aimSize,prname):
    '''
    HEADINGS
    '''
    pg1 = doc.add_paragraph()
    # Center pg1
    pg1.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    # Set Font to Times New Roman
    pg1.style.font.name = 'Times New Roman'
    pg1.style.font.bold = True
    pg1.style.font.size = aw.shared.Pt(headingSize)
    pg1.style.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    # Add heading text with specified font style
    pg1_run = pg1.add_run(heading)
    pg1_run.bold = True
    pg1_run.font.name = 'Times New Roman'
    pg1_run.font.size = aw.shared.Pt(headingSize)
    pg1_run.font.color.rgb = aw.shared.RGBColor(0, 0, 0)


    '''
    AIM
    '''
    pg2 = doc.add_paragraph()
    # Justify pg2
    pg2.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    # Set Font to Cambria(Body Text)
    pg2.style.font.name = 'Cambria'
    pg2.style.font.size = aw.shared.Pt(aimSize)
    pg2.style.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    # Add aim text with specified font style
    pg2_run1 = pg2.add_run('Aim: ')
    pg2_run1.bold = True
    pg2_run1.font.size = aw.shared.Pt(aimSize)
    pg2_run1.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    pg2_run2 = pg2.add_run(aim)
    pg2_run2.bold = True
    pg2_run2.font.size = aw.shared.Pt(aimSize)
    pg2_run2.font.color.rgb = aw.shared.RGBColor(0, 0, 0)


    # Read File pr01.py and add to document
    with open('pr01.py', 'r') as f:
        code = f.read()
    pg3 = doc.add_paragraph().add_run(code)
    pg3.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    # pg3_run = pg3.add_run(code)
    pg3.font.name = 'Cambria'
    pg3.font.size = aw.shared.Pt(aimSize)
    pg3.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    pg3.font.bold = False


    # Add Page Break
    doc.add_page_break()
    pg4 = doc.add_paragraph().add_run("OUTPUT:")
    pg4.alignment = aw.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    # pg4_run = pg4.add_run(code)
    pg4.font.name = 'Cambria'
    pg4.font.size = aw.shared.Pt(aimSize)
    pg4.font.color.rgb = aw.shared.RGBColor(0, 0, 0)
    pg4.font.bold = True





    # Run Python Files in new windows(Python File name will change according to heading)
    os.system(f'start cmd /k python {prname}.py')

    # Wait until program is executed and screenshot is taken
    time.sleep(2)





    pyautogui.hotkey('alt', 'enter')
    # max_x,max_y = pyautogui.center(pyautogui.locateOnScreen('maximize.png'))
    # print(max_x,max_y)
    # pyautogui.click(max_x,max_y)
    # time.sleep(1)

    pyautogui.screenshot('output.png', region=(0, 0, 1800, 880)).save('output.png')
    pyautogui.typewrite('exit')
    pyautogui.press('enter')

    # Add Screenshot to document
    doc.add_picture('output.png', width=aw.shared.Inches(5.5))

    doc.add_page_break()


# Add Heading and Aim to document
for i in range(len(aimList)):
    heading = f'Practical - {i+1}'
    makedoc(heading,aimList[i],headingSize,aimSize,f'pr0{i+1}')

# Save document
doc.save('test.docx')
# Open Document
os.system(f'start test.docx')
pyautogui.scroll(20)

